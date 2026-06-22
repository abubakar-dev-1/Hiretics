# -*- coding: utf-8 -*-
"""
Build a professional Word (.docx) report from FYP_FINAL_REPORT.md.
- Times New Roman 12 body, justified, 1.5 spacing
- H1 16 / H2 14 / H3 13 / H4 12 (bold, navy)
- Generated cover page + auto Table of Contents (Word field) + page numbers
- GitHub tables, fenced code blocks, lists, bold/italic/inline-code
- Embeds rendered diagrams at their figure captions
"""
import os
import re
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT = r"d:\Hiretics-new"
MD = os.path.join(ROOT, "FYP_FINAL_REPORT.md")
DIAG = os.path.join(ROOT, "report_assets", "diagrams")
OUT = os.path.join(ROOT, "Hiretics_FYP_Report.docx")

NAVY = RGBColor(0x1F, 0x3A, 0x5F)
BLACK = RGBColor(0x00, 0x00, 0x00)
LINK = RGBColor(0x1B, 0x57, 0x8C)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)

EN_DASH = "–"

FIG_MAP = {
    1: "fig1_usecase.png",
    2: "fig2_dfd0.png",
    3: "fig3_dfd1.png",
    4: "fig4_dfd2.png",
    5: "fig5_architecture.png",
    6: "fig6_class.png",
    7: "fig7_sequence.png",
    8: "fig8_sequence_stripe.png",
    9: "fig9_sequence_candidate.png",
    10: "fig10_collaboration.png",
    11: "fig11_erd.png",
    12: "fig12_deployment.png",
}


# ----------------------------------------------------------------------------
# styles / helpers
# ----------------------------------------------------------------------------
def set_cell_font(cell, size=10.5, bold=False, color=None):
    for p in cell.paragraphs:
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.line_spacing = 1.0
        for r in p.runs:
            r.font.name = "Times New Roman"
            r.font.size = Pt(size)
            r.font.bold = bold
            if color is not None:
                r.font.color.rgb = color


def shade(paragraph, fill="F2F2F2"):
    pPr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)
    pPr.append(shd)


def shade_cell(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)
    tcPr.append(shd)


def pbreak(doc):
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)


def setup_styles(doc):
    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(12)
    normal.font.color.rgb = BLACK
    rpr = normal.element.get_or_add_rPr()
    rfonts = rpr.get_or_add_rFonts()
    rfonts.set(qn("w:ascii"), "Times New Roman")
    rfonts.set(qn("w:hAnsi"), "Times New Roman")
    rfonts.set(qn("w:cs"), "Times New Roman")
    pf = normal.paragraph_format
    pf.line_spacing = 1.5
    pf.space_after = Pt(6)
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    for name, size in (("Heading 1", 16), ("Heading 2", 14), ("Heading 3", 13), ("Heading 4", 12)):
        st = doc.styles[name]
        st.font.name = "Times New Roman"
        st.font.size = Pt(size)
        st.font.bold = True
        st.font.color.rgb = NAVY
        st.paragraph_format.space_before = Pt(14 if size >= 14 else 10)
        st.paragraph_format.space_after = Pt(6)
        st.paragraph_format.keep_with_next = True
        st.paragraph_format.line_spacing = 1.2

    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)


# ----------------------------------------------------------------------------
# inline formatting:  **bold**  *italic*  `code`  [text](url)
# ----------------------------------------------------------------------------
INLINE = re.compile(r"(\*\*.+?\*\*|\*[^*]+?\*|`[^`]+?`|\[[^\]]+?\]\([^)]+?\))")


def add_runs(p, text, base_size=12):
    for tok in INLINE.split(text):
        if not tok:
            continue
        if tok.startswith("**") and tok.endswith("**"):
            r = p.add_run(tok[2:-2]); r.bold = True
        elif tok.startswith("*") and tok.endswith("*"):
            r = p.add_run(tok[1:-1]); r.italic = True
        elif tok.startswith("`") and tok.endswith("`"):
            r = p.add_run(tok[1:-1]); r.font.name = "Consolas"; r.font.size = Pt(base_size - 1.5)
        else:
            m = re.match(r"\[([^\]]+?)\]\(([^)]+?)\)", tok)
            if m:
                r = p.add_run(m.group(1)); r.font.color.rgb = LINK
            else:
                p.add_run(tok)


# ----------------------------------------------------------------------------
# Word fields (TOC + page numbers)
# ----------------------------------------------------------------------------
def add_field(paragraph, instr, placeholder):
    r1 = paragraph.add_run(); f1 = OxmlElement("w:fldChar"); f1.set(qn("w:fldCharType"), "begin"); r1._r.append(f1)
    r2 = paragraph.add_run(); it = OxmlElement("w:instrText"); it.set(qn("xml:space"), "preserve"); it.text = instr; r2._r.append(it)
    r3 = paragraph.add_run(); f3 = OxmlElement("w:fldChar"); f3.set(qn("w:fldCharType"), "separate"); r3._r.append(f3)
    paragraph.add_run(placeholder)
    r5 = paragraph.add_run(); f5 = OxmlElement("w:fldChar"); f5.set(qn("w:fldCharType"), "end"); r5._r.append(f5)


def add_toc(doc):
    h = doc.add_paragraph("Table of Contents"); h.style = doc.styles["Heading 1"]
    p = doc.add_paragraph()
    add_field(p, 'TOC \\o "1-3" \\h \\z \\u',
              "Right-click here and choose 'Update Field' (or press Ctrl+A then F9) to build the contents.")


def add_page_numbers(doc):
    footer = doc.sections[0].footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(); f = OxmlElement("w:fldChar"); f.set(qn("w:fldCharType"), "begin"); r._r.append(f)
    r = p.add_run(); it = OxmlElement("w:instrText"); it.set(qn("xml:space"), "preserve"); it.text = "PAGE"; r._r.append(it)
    r = p.add_run(); f = OxmlElement("w:fldChar"); f.set(qn("w:fldCharType"), "end"); r._r.append(f)
    for run in p.runs:
        run.font.name = "Times New Roman"; run.font.size = Pt(10)


# ----------------------------------------------------------------------------
# cover page
# ----------------------------------------------------------------------------
def cover(doc):
    def ctr(text, size, bold=True, color=NAVY, after=6, italic=False, before=0):
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(after); p.paragraph_format.space_before = Pt(before)
        p.paragraph_format.line_spacing = 1.15
        r = p.add_run(text); r.bold = bold; r.italic = italic
        r.font.size = Pt(size); r.font.color.rgb = color; r.font.name = "Times New Roman"
        return p

    ctr("University of Management and Technology", 16, color=NAVY, after=2, before=12)
    ctr("School of Systems & Technology", 13, bold=False, color=BLACK, after=2)
    ctr("Department of Informatics & Systems", 13, bold=False, color=BLACK, after=22)
    ctr("FINAL YEAR PROJECT REPORT", 14, color=BLACK, after=28)

    ctr("Hiretics", 30, color=NAVY, after=8)
    ctr("An Event-Driven, Queue-Based Serverless Architecture with "
        "Microservices-Style Domain Decomposition, Emulated Locally via LocalStack",
        14, bold=False, color=BLACK, italic=True, after=28)

    ctr("Submitted By", 12, color=BLACK, after=4)
    for nm, rn in (("Muhammad Abubakar", "F2022266276"),
                   ("Muhammad Ibrahim", "F2022266385"),
                   ("Eman Fatima", "F2022266766"),
                   ("Syeda Asfoora Iqbal", "F2022266102")):
        ctr("%s    (%s)" % (nm, rn), 12, bold=False, color=BLACK, after=2)

    ctr("Project Advisor:  Sir Hafiz Ahsan Arshad", 12, bold=False, color=BLACK, after=2, before=18)
    ctr("Co-Advisor:  Sir Awais Amin", 12, bold=False, color=BLACK, after=22)
    ctr("Session 2022 %s 2026" % EN_DASH, 12, color=BLACK, after=2)
    ctr("Lahore, Pakistan", 12, bold=False, color=BLACK, after=2)

    pbreak(doc)


# ----------------------------------------------------------------------------
# table parsing
# ----------------------------------------------------------------------------
SEP_RE = re.compile(r"^\s*\|?\s*:?-{2,}:?\s*(\|\s*:?-{2,}:?\s*)+\|?\s*$")


def split_row(line):
    line = line.strip()
    if line.startswith("|"):
        line = line[1:]
    if line.endswith("|"):
        line = line[:-1]
    return [c.strip() for c in line.split("|")]


def add_table(doc, rows):
    header = split_row(rows[0])
    body = [split_row(r) for r in rows[2:]]
    ncol = len(header)
    t = doc.add_table(rows=1, cols=ncol)
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, cell in enumerate(t.rows[0].cells):
        cell.paragraphs[0].text = ""
        add_runs(cell.paragraphs[0], header[j] if j < len(header) else "")
        set_cell_font(cell, size=10.5, bold=True, color=WHITE)
        shade_cell(cell, "1F3A5F")
    for brow in body:
        cells = t.add_row().cells
        for j in range(ncol):
            cells[j].paragraphs[0].text = ""
            add_runs(cells[j].paragraphs[0], brow[j] if j < len(brow) else "")
            set_cell_font(cells[j], size=10.5)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


# ----------------------------------------------------------------------------
# main conversion
# ----------------------------------------------------------------------------
def convert():
    raw = open(MD, encoding="utf-8").read()
    raw = re.sub(r"<!--.*?-->", "", raw, flags=re.DOTALL)        # strip HTML comments
    raw = re.sub(r"\*\*\(Render.*?\)\*\*", "", raw)              # strip "render yourself" notes
    lines = raw.split("\n")

    # content starts at "## Dedication"; the title block is replaced by the cover page
    start = next((i for i, l in enumerate(lines) if l.strip().lower().startswith("## dedication")), 0)
    lines = lines[start:]

    doc = Document()
    setup_styles(doc)
    add_page_numbers(doc)
    cover(doc)
    add_toc(doc)
    pbreak(doc)

    i, n = 0, len(lines)
    while i < n:
        line = lines[i]
        s = line.strip()

        if not s:
            i += 1
            continue

        # fenced code block
        if s.startswith("```"):
            i += 1
            code = []
            while i < n and not lines[i].strip().startswith("```"):
                code.append(lines[i]); i += 1
            i += 1
            p = doc.add_paragraph()
            p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.line_spacing = 1.0
            p.paragraph_format.space_after = Pt(8)
            shade(p, "F2F2F2")
            run = p.add_run("\n".join(code))
            run.font.name = "Consolas"; run.font.size = Pt(9.5)
            continue

        # horizontal rule -> skip
        if re.match(r"^-{3,}$", s) or re.match(r"^\*{3,}$", s):
            i += 1
            continue

        # table
        if "|" in s and i + 1 < n and SEP_RE.match(lines[i + 1]):
            block = [line]
            i += 1
            block.append(lines[i]); i += 1
            while i < n and "|" in lines[i] and lines[i].strip():
                block.append(lines[i]); i += 1
            add_table(doc, block)
            continue

        # headings
        m = re.match(r"^(#{1,6})\s+(.*)$", s)
        if m:
            level = len(m.group(1))
            text = re.sub(r"\*\*(.+?)\*\*", r"\1", m.group(2).strip())
            if level == 2:
                pbreak(doc)
            p = doc.add_paragraph(style="Heading %d" % min(level, 4))
            p.add_run(text)
            i += 1
            continue

        # figure caption -> insert diagram, then italic centered caption
        fm = re.match(r"^\*Figure\s+(\d+)", s)
        if fm:
            num = int(fm.group(1))
            img = FIG_MAP.get(num)
            if img and os.path.exists(os.path.join(DIAG, img)):
                ip = doc.add_paragraph(); ip.alignment = WD_ALIGN_PARAGRAPH.CENTER
                ip.paragraph_format.space_before = Pt(6); ip.paragraph_format.space_after = Pt(2)
                ip.add_run().add_picture(os.path.join(DIAG, img), width=Inches(6.2))
            cap = doc.add_paragraph(); cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cap.paragraph_format.space_after = Pt(10)
            r = cap.add_run(s.strip().strip("*")); r.italic = True; r.font.size = Pt(10)
            i += 1
            continue

        # bullet list
        if re.match(r"^[-*]\s+", s):
            p = doc.add_paragraph(style="List Bullet")
            p.paragraph_format.line_spacing = 1.3
            add_runs(p, re.sub(r"^[-*]\s+", "", s))
            i += 1
            continue

        # numbered list
        if re.match(r"^\d+\.\s+", s):
            p = doc.add_paragraph(style="List Number")
            p.paragraph_format.line_spacing = 1.3
            add_runs(p, re.sub(r"^\d+\.\s+", "", s))
            i += 1
            continue

        # blockquote
        if s.startswith(">"):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.4)
            add_runs(p, re.sub(r"^>\s?", "", s))
            for r in p.runs:
                r.italic = True
            i += 1
            continue

        # normal paragraph (gather wrapped lines)
        para = [s]
        i += 1
        while i < n:
            nxt = lines[i].strip()
            if (not nxt or nxt.startswith("#") or nxt.startswith("```") or nxt.startswith(">")
                    or re.match(r"^[-*]\s+", nxt) or re.match(r"^\d+\.\s+", nxt)
                    or re.match(r"^-{3,}$", nxt) or nxt.startswith("*Figure")
                    or ("|" in nxt and i + 1 < n and SEP_RE.match(lines[i + 1]))):
                break
            para.append(nxt)
            i += 1
        add_runs(doc.add_paragraph(), " ".join(para))

    out = OUT
    try:
        doc.save(out)
    except PermissionError:
        out = OUT.replace(".docx", "_NEW.docx")
        doc.save(out)
        print("NOTE: original was locked (open in Word). Saved a fresh copy instead.")
    print("SAVED ->", out)
    print("paragraphs:", len(doc.paragraphs), " tables:", len(doc.tables), " images:", len(doc.inline_shapes))


if __name__ == "__main__":
    convert()
